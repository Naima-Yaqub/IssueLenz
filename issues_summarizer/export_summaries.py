import io
import pandas as pd
from docx import Document
from fpdf import FPDF
import json
import csv
import streamlit as st
import os
import re
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import openpyxl
from openpyxl.styles import Font

def export_summaries(export_format, issues):
    if export_format != "None":
        try:
            if export_format == "Excel":
                xlsx_data = export_to_excel(issues)
                st.download_button("Download Excel File", xlsx_data, "issues_summary.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
            elif export_format == "Word":
                word_data = export_to_word(issues)
                st.download_button(
                    "Download Word Document",
                    word_data,
                    "issues_summary.docx",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            elif export_format == "JSON":
                json_data = export_to_json(issues)
                st.download_button("Download JSON File", json_data, "issues_summary.json", "application/json")
        except Exception as e:
            st.error(f"An error occurred while exporting the file: {str(e)}")

def export_to_excel(issues):
    if not issues:
        return None

    wb = openpyxl.Workbook()
    ws = wb.active

    headers = ["number", "title", "description", "created_at", "labels", "url"]
    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_num, value=header)
        cell.font = Font(bold=True)

    for row_num, issue in enumerate(issues, 2):
        for col_num, field in enumerate(headers, 1):
            cell_value = issue.get(field, "")
            if field == "labels":
                cell_value = "No Label" if not cell_value else ", ".join(cell_value)
            ws.cell(row=row_num, column=col_num, value=cell_value)

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)

    return output

def export_to_word(issues):
    doc = Document()
    for issue in issues:
        doc.add_heading(f"Issue #{issue['number']}: {issue['title']}", level=1)
        doc.add_paragraph(f"Description: {issue['description']}")
        doc.add_paragraph(f"Created At: {issue['created_at']}")
        labels = issue.get('labels', ['No label'])
        doc.add_paragraph(f"Labels: {', '.join(labels)}")
        doc.add_paragraph(f"URL: {issue['url']}")
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def export_to_json(issues):
    for issue in issues:
        if not issue.get('labels'):
            issue['labels'] = ['No label']
    return json.dumps(issues, indent=4).encode('utf-8')
